# ══════════════════════════════════════════════════════
# ACADEMIC RESEARCH ASSISTANT — WEB VERSION
# Author: Mahdi Mahdavi PhD
# University of Roehampton
# Run with: streamlit run app.py
# ══════════════════════════════════════════════════════

import streamlit as st
import anthropic
import requests
import pandas as pd
import os
import json
import io
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

# ── PAGE CONFIG ───────────────────────────────────────
st.set_page_config(
    page_title="Academic Research Assistant",
    page_icon="🎓",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ── CUSTOM CSS ────────────────────────────────────────
st.markdown("""
<style>
    /* Main background */
    .main {
        background-color: #0f1117;
    }

    /* Sidebar */
    .css-1d391kg {
        background-color: #1a1d2e;
    }

    /* Headers */
    h1, h2, h3 {
        color: #4da6ff;
    }

    /* Cards */
    .result-card {
        background-color: #1e2130;
        border-radius: 10px;
        padding: 20px;
        margin: 10px 0;
        border-left: 4px solid #4da6ff;
    }

    /* Success box */
    .success-box {
        background-color: #1a2e1a;
        border-radius: 8px;
        padding: 15px;
        border-left: 4px solid #00cc44;
        color: #00cc44;
    }

    /* Info box */
    .info-box {
        background-color: #1a1f2e;
        border-radius: 8px;
        padding: 15px;
        border-left: 4px solid #4da6ff;
        color: #4da6ff;
    }

    /* Metric cards */
    .metric-card {
        background-color: #1e2130;
        border-radius: 10px;
        padding: 15px;
        text-align: center;
        border: 1px solid #2d3250;
    }

    /* Button styling */
    .stButton > button {
        background-color: #4da6ff;
        color: white;
        border-radius: 8px;
        border: none;
        padding: 10px 20px;
        font-weight: bold;
        width: 100%;
    }

    .stButton > button:hover {
        background-color: #2d8af0;
    }

    /* Download button */
    .stDownloadButton > button {
        background-color: #00cc44;
        color: white;
        border-radius: 8px;
        border: none;
    }

    /* Text areas */
    .stTextArea textarea {
        background-color: #1e2130;
        color: #ffffff;
        border: 1px solid #2d3250;
        border-radius: 8px;
    }

    /* Input fields */
    .stTextInput input {
        background-color: #1e2130;
        color: #ffffff;
        border: 1px solid #2d3250;
        border-radius: 8px;
    }

    /* Select boxes */
    .stSelectbox select {
        background-color: #1e2130;
        color: #ffffff;
    }

    /* Hide streamlit branding */
    #MainMenu {visibility: hidden;}
    footer {visibility: hidden;}
    header {visibility: hidden;}
</style>
""", unsafe_allow_html=True)

# ── API SETUP ─────────────────────────────────────────
@st.cache_resource
def get_client():
    return anthropic.Anthropic(
        api_key=st.secrets["ANTHROPIC_API_KEY"]
    )

client = get_client()

# ── SESSION STATE ─────────────────────────────────────
if "history" not in st.session_state:
    st.session_state.history = []
if "last_result" not in st.session_state:
    st.session_state.last_result = ""


# ══════════════════════════════════════════════════════
# CORE FUNCTIONS
# ══════════════════════════════════════════════════════

def ask_claude(prompt, system=None):
    """Send prompt to Claude"""
    if system is None:
        system = (
            "You are a personal research assistant "
            "for Mahdi Mahdavi PhD, Senior Lecturer "
            "in Health Services Management at the "
            "University of Roehampton. You specialise "
            "in health economics, NHS policy, HTA "
            "methodology and prevention economics. "
            "Write clearly with markdown formatting."
        )
    response = client.messages.create(
        model="claude-opus-4-5",
        max_tokens=4096,
        system=system,
        messages=[{
            "role": "user",
            "content": prompt
        }]
    )
    return response.content[0].text


def search_pubmed(query, max_results=10):
    """Search PubMed"""
    base = "https://eutils.ncbi.nlm.nih.gov/entrez/eutils/"
    try:
        search = requests.get(
            f"{base}esearch.fcgi",
            params={
                "db": "pubmed",
                "term": query,
                "retmax": max_results,
                "retmode": "json"
            }, timeout=10
        ).json()

        ids = search["esearchresult"]["idlist"]
        if not ids:
            return "", 0

        abstracts = requests.get(
            f"{base}efetch.fcgi",
            params={
                "db": "pubmed",
                "id": ",".join(ids),
                "retmode": "text",
                "rettype": "abstract"
            }, timeout=10
        ).text

        return abstracts, len(ids)
    except:
        return "", 0


def search_semantic_scholar(query, max_results=10):
    """Search Semantic Scholar"""
    try:
        response = requests.get(
            "https://api.semanticscholar.org/graph/v1/paper/search",
            params={
                "query": query,
                "limit": max_results,
                "fields": "title,abstract,authors,year,citationCount,url"
            }, timeout=10
        )
        papers = response.json().get("data", [])
        results = []
        for p in papers:
            results.append({
                "Title": p.get("title", "")[:60],
                "Authors": ", ".join([
                    a["name"]
                    for a in p.get("authors", [])[:2]
                ]),
                "Year": p.get("year", ""),
                "Citations": p.get("citationCount", 0),
                "Abstract": p.get("abstract", "")
            })
        return results
    except:
        return []


def search_openalex(query, max_results=10):
    """Search OpenAlex"""
    try:
        response = requests.get(
            "https://api.openalex.org/works",
            params={
                "search": query,
                "per-page": max_results,
                "select": "title,authorships,publication_year,cited_by_count,doi"
            },
            headers={"User-Agent": "mailto:m.mahdavi@roehampton.ac.uk"},
            timeout=10
        )
        papers = response.json().get("results", [])
        results = []
        for p in papers:
            results.append({
                "Title": (p.get("title") or "")[:60],
                "Authors": ", ".join([
                    a["author"]["display_name"]
                    for a in p.get("authorships", [])[:2]
                ]),
                "Year": p.get("publication_year", ""),
                "Citations": p.get("cited_by_count", 0),
                "DOI": p.get("doi", "")
            })
        return results
    except:
        return []


def create_word_bytes(title, content, author):
    """Create Word document and return as bytes"""
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin   = Inches(1.2)
        section.right_margin  = Inches(1.2)

    t = doc.add_heading(title, 0)
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER

    m = doc.add_paragraph(
        f"Author: {author}  |  "
        f"University of Roehampton  |  "
        f"{date.today().strftime('%d %B %Y')}"
    )
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    m.runs[0].font.size = Pt(9)
    doc.add_paragraph()

    for line in content.split("\n"):
        line = line.strip()
        if not line:
            doc.add_paragraph()
        elif line.startswith("## "):
            doc.add_heading(
                line.replace("## ", ""), 1)
        elif line.startswith("### "):
            doc.add_heading(
                line.replace("### ", ""), 2)
        elif line.startswith("# "):
            doc.add_heading(
                line.replace("# ", ""), 1)
        elif (line.startswith("- ") or
              line.startswith("→ ") or
              line.startswith("• ")):
            p = doc.add_paragraph(
                style="List Bullet")
            p.add_run(
                line.lstrip("-→• ").strip())
        else:
            doc.add_paragraph(line)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def create_excel_bytes(data, title=None):
    """Create Excel file and return as bytes"""
    if isinstance(data, dict):
        df = pd.DataFrame(data)
    else:
        df = data.copy()

    buf = io.BytesIO()
    wb  = openpyxl.Workbook()
    ws  = wb.active

    start_row = 1
    if title:
        cell = ws.cell(row=1, column=1, value=title)
        cell.font = Font(bold=True, size=14)
        ws.merge_cells(
            start_row=1, start_column=1,
            end_row=1, end_column=len(df.columns)
        )
        start_row = 3

    hfill = PatternFill(
        fill_type="solid", fgColor="2166AC")
    for i, col in enumerate(df.columns, 1):
        c = ws.cell(
            row=start_row, column=i, value=str(col))
        c.font  = Font(bold=True, color="FFFFFF")
        c.fill  = hfill
        c.alignment = Alignment(horizontal="center")

    lfill = PatternFill(
        fill_type="solid", fgColor="EBF5FB")
    for ri, row in enumerate(
        df.itertuples(index=False),
        start=start_row + 1
    ):
        for ci, val in enumerate(row, 1):
            c = ws.cell(row=ri, column=ci, value=val)
            if ri % 2 == 0:
                c.fill = lfill

    for col in ws.columns:
        w = max(
            (len(str(c.value)) for c in col if c.value),
            default=10
        )
        ws.column_dimensions[
            col[0].column_letter
        ].width = min(w + 4, 45)

    wb.save(buf)
    buf.seek(0)
    return buf.getvalue()


# ══════════════════════════════════════════════════════
# SIDEBAR
# ══════════════════════════════════════════════════════

with st.sidebar:
    st.image(
        "https://upload.wikimedia.org/wikipedia/en/thumb/e/ef/Roehampton_University_logo.svg/220px-Roehampton_University_logo.svg.png",
        width=150
    )

    st.markdown("---")
    st.markdown(
        "**Mahdi Mahdavi PhD**\n\n"
        "Senior Lecturer\n\n"
        "Health Services Management\n\n"
        "University of Roehampton"
    )
    st.markdown("---")

    page = st.selectbox(
        "Navigation",
        options=[
            "🏠 Home",
            "📚 Literature Review",
            "📊 Data Analysis",
            "✍️ Write Section",
            "📧 Draft Email",
            "❓ Ask Anything",
            "📋 History"
        ]
    )

    st.markdown("---")
    st.markdown(
        f"*{date.today().strftime('%d %B %Y')}*")


# ══════════════════════════════════════════════════════
# PAGES
# ══════════════════════════════════════════════════════

# ── HOME PAGE ─────────────────────────────────────────
if page == "🏠 Home":

    st.title("🎓 Academic Research Assistant")
    st.markdown(
        "*Personal AI assistant for health economics research*"
    )

    st.markdown("---")

    col1, col2, col3 = st.columns(3)

    with col1:
        st.markdown("""
        <div class='metric-card'>
        <h3>📚</h3>
        <h4>Literature Review</h4>
        <p>Search PubMed, Semantic Scholar
        and OpenAlex simultaneously</p>
        </div>
        """, unsafe_allow_html=True)

    with col2:
        st.markdown("""
        <div class='metric-card'>
        <h3>📊</h3>
        <h4>Data Analysis</h4>
        <p>Analyse any dataset and generate
        Word and Excel reports</p>
        </div>
        """, unsafe_allow_html=True)

    with col3:
        st.markdown("""
        <div class='metric-card'>
        <h3>✍️</h3>
        <h4>Academic Writing</h4>
        <p>Write methods, discussion,
        grant sections and reports</p>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")

    col4, col5, col6 = st.columns(3)

    with col4:
        st.markdown("""
        <div class='metric-card'>
        <h3>📧</h3>
        <h4>Email Drafting</h4>
        <p>Professional emails for
        clients, colleagues and journals</p>
        </div>
        """, unsafe_allow_html=True)

    with col5:
        st.markdown("""
        <div class='metric-card'>
        <h3>❓</h3>
        <h4>Ask Anything</h4>
        <p>Health economics, NHS policy,
        NICE methodology, HTA questions</p>
        </div>
        """, unsafe_allow_html=True)

    with col6:
        st.markdown("""
        <div class='metric-card'>
        <h3>📋</h3>
        <h4>History</h4>
        <p>Review all previous tasks
        and download saved files</p>
        </div>
        """, unsafe_allow_html=True)

    st.markdown("---")
    st.info(
        "Select a tool from the sidebar to get started."
    )


# ── LITERATURE REVIEW PAGE ────────────────────────────
elif page == "📚 Literature Review":

    st.title("📚 Literature Review")
    st.markdown(
        "Search multiple academic databases "
        "simultaneously and synthesise the evidence."
    )

    with st.form("lit_review_form"):
        topic = st.text_input(
            "Research Topic",
            placeholder="e.g. Cost-effectiveness of adult vaccination in England"
        )

        col1, col2 = st.columns(2)
        with col1:
            pubmed_query = st.text_input(
                "PubMed Search Query",
                placeholder="e.g. vaccination cost effectiveness elderly England"
            )
            openalex_query = st.text_input(
                "OpenAlex Search Query",
                placeholder="e.g. prevention economics NHS"
            )
        with col2:
            scholar_query = st.text_input(
                "Semantic Scholar Query",
                placeholder="e.g. adult immunisation economic returns UK"
            )
            wos_file = st.file_uploader(
                "Web of Science Export (CSV)",
                type=["csv"]
            )

        max_results = st.slider(
            "Results per database", 5, 20, 10)

        filename = st.text_input(
            "Output filename",
            value="literature_review"
        )

        submitted = st.form_submit_button(
            "🔍 Search and Synthesise",
            use_container_width=True
        )

    if submitted and topic:
        all_results = []
        papers_data = []

        # Search databases
        if pubmed_query:
            with st.spinner("Searching PubMed..."):
                r, n = search_pubmed(
                    pubmed_query, max_results)
                if r:
                    all_results.append(
                        f"=== PUBMED ({n} papers) ===\n{r}")
                    st.success(
                        f"PubMed: {n} papers found")

        if scholar_query:
            with st.spinner(
                "Searching Semantic Scholar..."
            ):
                papers = search_semantic_scholar(
                    scholar_query, max_results)
                if papers:
                    papers_data.extend(papers)
                    r = "\n".join([
                        f"TITLE: {p['Title']}\n"
                        f"AUTHORS: {p['Authors']}\n"
                        f"YEAR: {p['Year']}\n"
                        f"ABSTRACT: {p['Abstract']}"
                        for p in papers
                    ])
                    all_results.append(
                        f"=== SEMANTIC SCHOLAR ===\n{r}")
                    st.success(
                        f"Semantic Scholar: "
                        f"{len(papers)} papers found"
                    )

        if openalex_query:
            with st.spinner("Searching OpenAlex..."):
                papers = search_openalex(
                    openalex_query, max_results)
                if papers:
                    papers_data.extend(papers)
                    r = "\n".join([
                        f"TITLE: {p['Title']}\n"
                        f"AUTHORS: {p['Authors']}\n"
                        f"YEAR: {p['Year']}"
                        for p in papers
                    ])
                    all_results.append(
                        f"=== OPENALEX ===\n{r}")
                    st.success(
                        f"OpenAlex: "
                        f"{len(papers)} papers found"
                    )

        if wos_file:
            df_wos = pd.read_csv(
                wos_file, encoding="utf-8-sig")
            r = "\n".join([
                f"TITLE: {row.get('Article Title','')}\n"
                f"AUTHORS: {row.get('Authors','')}\n"
                f"YEAR: {row.get('Publication Year','')}\n"
                f"ABSTRACT: {row.get('Abstract','')}"
                for _, row in df_wos.iterrows()
            ])
            all_results.append(
                f"=== WEB OF SCIENCE ===\n{r}")
            st.success(
                f"WoS: {len(df_wos)} papers loaded")

        if papers_data:
            st.subheader("Papers Found")
            df_papers = pd.DataFrame(papers_data)
            st.dataframe(
                df_papers,
                use_container_width=True
            )

        if all_results:
            combined = "\n\n".join(all_results)
            with st.spinner(
                "Claude is synthesising the literature..."
            ):
                review = ask_claude(f"""
Conduct a structured literature review on:
{topic}

Search results:
{combined[:8000]}

Write with these sections:
## 1. Overview
## 2. Key Themes
## 3. Most Important Studies
## 4. Evidence Quality
## 5. Evidence Gaps
## 6. UK and NHS Implications
## 7. Key References (Vancouver format)

Academic style for UK health economics.
""")

            st.subheader("Literature Review")
            st.markdown(review)

            st.session_state.history.append({
                "type": "Literature Review",
                "topic": topic,
                "result": review,
                "date": str(date.today())
            })
            st.session_state.last_result = review

            # Download buttons
            col1, col2 = st.columns(2)
            with col1:
                word_bytes = create_word_bytes(
                    title=f"Literature Review: {topic}",
                    content=review,
                    author="Mahdi Mahdavi PhD"
                )
                st.download_button(
                    "📄 Download Word Report",
                    data=word_bytes,
                    file_name=f"{filename}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True
                )

            with col2:
                if papers_data:
                    excel_bytes = create_excel_bytes(
                        data=pd.DataFrame(papers_data),
                        title=f"Papers: {topic}"
                    )
                    st.download_button(
                        "📊 Download Papers Excel",
                        data=excel_bytes,
                        file_name=f"{filename}_papers.xlsx",
                        mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                        use_container_width=True
                    )


# ── DATA ANALYSIS PAGE ────────────────────────────────
elif page == "📊 Data Analysis":

    st.title("📊 Data Analysis")
    st.markdown(
        "Analyse your data and generate "
        "professional reports."
    )

    with st.form("data_form"):
        task = st.text_input(
            "What do you want to analyse?",
            placeholder="e.g. Vaccination coverage gaps across age groups"
        )

        input_method = st.radio(
            "Data input method",
            ["Upload Excel/CSV file",
             "Enter data manually",
             "No data — just analyse from description"]
        )

        uploaded_file = None
        manual_data  = None

        if input_method == "Upload Excel/CSV file":
            uploaded_file = st.file_uploader(
                "Upload your data file",
                type=["xlsx", "csv", "xls"]
            )

        elif input_method == "Enter data manually":
            manual_data = st.text_area(
                "Paste your data (JSON or CSV format)",
                height=150,
                placeholder='{"Vaccine": ["Flu", "RSV"], "Coverage": [74.9, 61.9]}'
            )

        filename = st.text_input(
            "Output filename",
            value="analysis"
        )

        submitted = st.form_submit_button(
            "📊 Analyse",
            use_container_width=True
        )

    if submitted and task:
        data = None
        data_str = ""

        if uploaded_file:
            with st.spinner("Reading file..."):
                if uploaded_file.name.endswith(".csv"):
                    data = pd.read_csv(uploaded_file)
                else:
                    data = pd.read_excel(uploaded_file)
                st.success(
                    f"Loaded {len(data)} rows")

        elif manual_data:
            try:
                data = json.loads(manual_data)
                data = pd.DataFrame(data)
                st.success("Data loaded")
            except:
                st.error("Could not parse data. "
                         "Check JSON format.")

        if data is not None:
            st.subheader("Your Data")
            st.dataframe(
                data, use_container_width=True)
            data_str = data.to_string()

        with st.spinner(
            "Claude is analysing..."
        ):
            analysis = ask_claude(f"""
Task: {task}
Data: {data_str[:4000]}

Provide structured analysis:
## 1. Summary
## 2. Main Findings
## 3. Key Insights
## 4. NHS and Policy Implications
## 5. Limitations
## 6. Recommendations
""")

        st.subheader("Analysis")
        st.markdown(analysis)

        st.session_state.history.append({
            "type": "Data Analysis",
            "topic": task,
            "result": analysis,
            "date": str(date.today())
        })

        col1, col2 = st.columns(2)
        with col1:
            word_bytes = create_word_bytes(
                title=task,
                content=analysis,
                author="Mahdi Mahdavi PhD"
            )
            st.download_button(
                "📄 Download Word Report",
                data=word_bytes,
                file_name=f"{filename}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )

        with col2:
            if data is not None:
                excel_bytes = create_excel_bytes(
                    data=data,
                    title=task
                )
                st.download_button(
                    "📊 Download Excel",
                    data=excel_bytes,
                    file_name=f"{filename}.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                    use_container_width=True
                )


# ── WRITE SECTION PAGE ────────────────────────────────
elif page == "✍️ Write Section":

    st.title("✍️ Write Academic Section")
    st.markdown(
        "Generate any academic section instantly."
    )

    with st.form("write_form"):
        section_type = st.selectbox(
            "Section Type",
            options=[
                "Methods section",
                "Introduction",
                "Discussion",
                "Results narrative",
                "Executive summary",
                "Literature review section",
                "Grant application section",
                "Policy brief",
                "Abstract",
                "Conclusion",
                "Custom section"
            ]
        )

        if section_type == "Custom section":
            section_type = st.text_input(
                "Describe the section"
            )

        content_notes = st.text_area(
            "Key content to include",
            height=200,
            placeholder=(
                "Describe what should be in this section...\n"
                "e.g. Three intervention areas: vaccination,\n"
                "integration, equity. ROI methodology.\n"
                "Green Book 2022. NICE PMG36."
            )
        )

        col1, col2 = st.columns(2)
        with col1:
            word_count = st.selectbox(
                "Target length",
                ["Short (200-300 words)",
                 "Medium (400-600 words)",
                 "Long (700-1000 words)",
                 "Very long (1000+ words)"]
            )
        with col2:
            audience = st.selectbox(
                "Target audience",
                ["Academic journal",
                 "Policy makers / MPs",
                 "NHS commissioners",
                 "Grant committee",
                 "General public"]
            )

        filename = st.text_input(
            "Output filename",
            value="section"
        )

        submitted = st.form_submit_button(
            "✍️ Write Section",
            use_container_width=True
        )

    if submitted and content_notes:
        with st.spinner(
            f"Writing {section_type}..."
        ):
            content = ask_claude(f"""
Write a {section_type} for {audience}.
Target length: {word_count}

Content to include:
{content_notes}

Requirements:
→ UK NHS context
→ NICE PMG36 and Green Book 2022
  where relevant
→ Professional academic tone
→ Clear subheadings
→ British English
""")

        st.subheader(section_type)
        st.markdown(content)

        st.session_state.history.append({
            "type": "Written Section",
            "topic": section_type,
            "result": content,
            "date": str(date.today())
        })

        word_bytes = create_word_bytes(
            title=section_type,
            content=content,
            author="Mahdi Mahdavi PhD"
        )
        st.download_button(
            "📄 Download Word Document",
            data=word_bytes,
            file_name=f"{filename}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )


# ── DRAFT EMAIL PAGE ──────────────────────────────────
elif page == "📧 Draft Email":

    st.title("📧 Draft Email")
    st.markdown(
        "Generate professional emails instantly."
    )

    with st.form("email_form"):
        situation = st.text_area(
            "Describe the email situation",
            height=150,
            placeholder=(
                "e.g. Following up with Hannah at MHP Group\n"
                "about a prevention economics project for GSK.\n"
                "She said they are engaging with client\n"
                "about supplier options."
            )
        )

        col1, col2 = st.columns(2)
        with col1:
            tone = st.selectbox(
                "Tone",
                ["Professional",
                 "Formal",
                 "Warm and friendly",
                 "Academic",
                 "Assertive"]
            )
        with col2:
            length = st.selectbox(
                "Length",
                ["Very short (under 100 words)",
                 "Short (100-150 words)",
                 "Medium (150-250 words)",
                 "Long (250+ words)"]
            )

        submitted = st.form_submit_button(
            "📧 Draft Email",
            use_container_width=True
        )

    if submitted and situation:
        with st.spinner("Drafting email..."):
            email = ask_claude(f"""
Draft a {tone.lower()} email.
Length: {length}
Situation: {situation}

Requirements:
→ British English
→ Include subject line
→ Professional greeting and sign off
→ Ready to send
→ Mahdi Mahdavi PhD, University of Roehampton
""")

        st.subheader("Drafted Email")
        st.markdown(email)

        st.session_state.history.append({
            "type": "Email Draft",
            "topic": situation[:50],
            "result": email,
            "date": str(date.today())
        })

        word_bytes = create_word_bytes(
            title="Email Draft",
            content=email,
            author="Mahdi Mahdavi PhD"
        )
        st.download_button(
            "📄 Download as Word",
            data=word_bytes,
            file_name="email_draft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )


# ── ASK ANYTHING PAGE ─────────────────────────────────
elif page == "❓ Ask Anything":

    st.title("❓ Ask Anything")
    st.markdown(
        "Ask any question about health economics, "
        "NHS policy, HTA or academic research."
    )

    # Quick question buttons
    st.subheader("Quick Questions")
    col1, col2, col3 = st.columns(3)

    with col1:
        if st.button("NICE QALY threshold 2024",
                     use_container_width=True):
            st.session_state.quick_q = (
                "What is the current NICE "
                "cost-effectiveness threshold "
                "for QALYs in England in 2024?"
            )
    with col2:
        if st.button("Green Book discount rate",
                     use_container_width=True):
            st.session_state.quick_q = (
                "What is the standard discount "
                "rate for health economic "
                "appraisal in the UK per the "
                "HM Treasury Green Book 2022?"
            )
    with col3:
        if st.button("Darzi Review key findings",
                     use_container_width=True):
            st.session_state.quick_q = (
                "What are the key economic "
                "findings of the Darzi Review "
                "2024 regarding NHS prevention?"
            )

    st.markdown("---")

    with st.form("ask_form"):
        default_q = getattr(
            st.session_state, "quick_q", "")

        question = st.text_area(
            "Your Question",
            value=default_q,
            height=120,
            placeholder="Ask anything about health economics, NHS, HTA..."
        )

        submitted = st.form_submit_button(
            "🔍 Ask Claude",
            use_container_width=True
        )

    if submitted and question:
        with st.spinner("Claude is answering..."):
            answer = ask_claude(question)

        st.subheader("Answer")
        st.markdown(answer)

        st.session_state.history.append({
            "type": "Question",
            "topic": question[:50],
            "result": answer,
            "date": str(date.today())
        })

        word_bytes = create_word_bytes(
            title=f"Q: {question[:50]}",
            content=f"**Question:**\n{question}\n\n**Answer:**\n{answer}",
            author="Mahdi Mahdavi PhD"
        )
        st.download_button(
            "📄 Download Answer",
            data=word_bytes,
            file_name="answer.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True
        )


# ── HISTORY PAGE ──────────────────────────────────────
elif page == "📋 History":

    st.title("📋 Session History")
    st.markdown(
        "All tasks completed in this session."
    )

    if not st.session_state.history:
        st.info(
            "No tasks completed yet. "
            "Use the tools in the sidebar."
        )
    else:
        for i, item in enumerate(
            reversed(st.session_state.history)
        ):
            with st.expander(
                f"{item['type']}: "
                f"{item['topic'][:50]} "
                f"— {item['date']}"
            ):
                st.markdown(item["result"])

                word_bytes = create_word_bytes(
                    title=item["topic"],
                    content=item["result"],
                    author="Mahdi Mahdavi PhD"
                )
                st.download_button(
                    "📄 Download",
                    data=word_bytes,
                    file_name=f"task_{i}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    key=f"dl_{i}"
                )
